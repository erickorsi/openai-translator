# -*- coding: utf-8 -*-
"""
Created on Tue May 09 11:52 2023

@author: erickorsi

Translate a document from one language to another.

v. 0.0.1

"""
import os
import time
from pathlib import Path
from enum import Enum

import openai
from docx import Document

openai.api_key = os.getenv("OPENAI_API_KEY")
#----------------------------------------
def translate(doc_path: Path,
              txt_path: Path,
              translated_path: Path,
              origin_lang: str = "en",
              target_lang: str = "pt"
              ) -> str:
    '''
    Translate a document from one language to another.

    Available languages:
        "en" - English
        "pt" - Portuguese
    '''
    # Get the text from the document
    full_text = __get_doc_text(doc_path, txt_path)

    origin_lang = Languages[origin_lang.upper()].value
    target_lang = Languages[target_lang.upper()].value

    # Translate the text
    prompt_system = f"""
The following text is a document in {origin_lang}.

Your task is to translate it to {target_lang}.

While translating, consider the associations between the words and the context of the text.
Maintain the same meaning and style of the overall text.
Maintain the same structure of the text, so as to compare it with the original text.
Maintain the paragraph breaks, and the tables structure, if tehy exist.

Return only the translated text.
    """

    completion_text = []
    for text in full_text:
        prompt_user = f"""
Text to be translated:

'''
{text}
'''
        """

        tries = 0
        while tries < 10:
            try:
                completion = __get_completion(prompt_system, prompt_user)
                print(completion)
                completion_text.append(completion)
                break
            except openai.error.RateLimitError:
                print("Rate limit reached. Waiting 1 minute...")
                time.sleep(65)
                tries += 1

    # Save the translated text to a txt file
    with open(translated_path, "w", encoding="utf-8") as txt_file:
        txt_file.write(completion_text[0])


def __get_doc_text(doc_path: Path,
                   txt_path: Path) -> str:
    '''
    Get the text from a document.
    Saves the text to a txt file.
    '''
    with open(doc_path, "rb") as doc_file:
        doc = Document(doc_file)
        # For reading over connection or streaming
        #source_stream = StringIO(f.read())
    #doc = Document(source_stream)

    full_text = []
    for para in doc.paragraphs:
        if para.text != "":
            full_text.append(para.text)

    full_text2 = []
    full_text2.append("\n\n".join(full_text))

    # Save the text to a txt file
    with open(txt_path, "w", encoding="utf-8") as txt_file:
        txt_file.write(full_text2[0])

    return full_text2


def __get_completion(prompt_system, prompt_user, model="gpt-3.5-turbo"):
    '''
    Get the completion of a prompt from the OpenAI API.
    '''
    messages = [{"role": "system", "content": prompt_system},
                {"role": "user", "content": prompt_user},]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0.8,
    )
    return response.choices[0].message["content"]


class Languages(Enum):
    '''
    Languages available for translation.
    '''
    PT = "portuguese"
    EN = "english"

if __name__ == "__main__":
    doc_path = Path(r"C:\Users\erick\OneDrive\Documentos\GitHub\openai-translator\examples\191056 - Mile High Cycles.docx")
    txtname = doc_path.stem + ".txt"
    txt_path = Path(doc_path.parent.parent / "txt-tests" / txtname)
    translated_path = Path(doc_path.parent.parent / "response-tests" / txtname)
    origin_lang = "en"
    target_lang = "pt"

    translate(doc_path, txt_path, translated_path, origin_lang, target_lang)
